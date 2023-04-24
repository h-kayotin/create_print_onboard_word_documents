"""
fts_onboard - FTS的入职单生成

Author: JiangHai江海
Date： 2023/4/24
"""
import openpyxl
from openpyxl.workbook.workbook import Workbook
from openpyxl.worksheet.worksheet import Worksheet
from docx import Document
from docx.document import Document as Doc
import docxcompose
from docxcompose.composer import Composer
import os


def read_from_excel():
    """
    从入职信息excel读取信息
    :return: 返回得到的员工信息列表
    """
    # 从入职信息读取workbook，再获取sheet
    wb = openpyxl.load_workbook("resources/入职信息-FTS.xlsx", data_only=True)  # type: Workbook
    sheet = wb["Sheet1"]  # type:Worksheet
    employees = []
    for row_num in range(2, sheet.max_row+1):
        employee = {
            "id": sheet[f"G{row_num}"].value,  # 工号
            "name_full": sheet[f"A{row_num}"].value,  # 姓名
            "DEPT": sheet[f"K{row_num}"].value,  # 部门
            "PC": sheet[f"O{row_num}"].value,  # 电脑编号
            "ad_account": sheet[f"F{row_num}"].value,  # 域账号
            "ad_pwd": sheet[f"I{row_num}"].value,  # 域账号初始密码
            "title": sheet[f"J{row_num}"].value,  # 职位
            "vpn_account": sheet[f"L{row_num}"].value,  # VPN账号
            "mail_account": sheet[f"M{row_num}"].value,  # 邮箱账号&wifi账号
            "vpn_pwd": sheet[f"N{row_num}"].value,  # VPN密码&WIFI密码
        }
        employees.append(employee)
    return employees


def read_word_temp():
    """
    读取word模板
    :return: 返回Document对象
    """
    doc = Document("resources/入职单模板FTS.docx")
    return doc


def make_onboard_words(employees):
    """
    根据员工数量生成n张入职单
    :param employees: 员工信息的列表
    :return: 成功返回True
    """
    for employee in employees:
        doc = read_word_temp()  # type:Doc
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in employee.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
        # 替换文本框中的占位符
        children = doc.element.body.iter()
        for child in children:
            if child.tag.endswith('txbx'):  # 获取所有文本框的tag
                for ci in child.iter():
                    if ci.tag.endswith('main}t'):  # 获取所有的行
                        for key, value in employee.items():
                            if key in ci.text:
                                ci.text = ci.text.replace(key, str(value))
        doc.save(f"onboard_words/fts-入职单{employee['vpn_account']}.docx")
    return True


def get_files_list(source_path):
    """
    获取指定文件路径的文件名列表
    :param source_path: 文件夹路径
    :return: 文件名列表
    """
    source_files_list = os.listdir(source_path)
    source_files = []
    for file in source_files_list:
        source_files.append(source_path + file)
    return source_files


def make_words_one(source_files, target_file):
    """
    将多个word合并成一个
    :param source_files: 源文件列表
    :param target_file: 合并的目标文件路径+名称
    :return: 成功返回True
    """
    # 新建一个文档对象作为第一页
    target_doc_first = Document(source_files[0])  # type: Doc
    # target_doc_first.add_page_break()
    target_composer = Composer(target_doc_first)

    # 从第二个文档开始循环，加入到新文档中
    for file in source_files[1:]:
        page_new = Document(file)  # type:Doc
        # page_new.add_page_break()
        target_composer.append(page_new)

    # 保存新文档
    target_composer.save(target_file)
    return True


def onboard_main():
    # 从入职信息excel读取信息
    employees = read_from_excel()
    print(f"获取到{len(employees)}条入职信息--->")
    # 根据入职信息生成N份入职表单
    if make_onboard_words(employees):
        print("生成入职单成功--->")
    source_path = "onboard_words/"  # 源文件路径
    target_file = "resources/入职单打印.docx"  # 目标路径
    # 获取源文件夹的文件列表，合并为一个新文档
    make_words_one(get_files_list(source_path), target_file)
    print(f"合并成功---^_^")
    input("Press Enter to quit--->")


if __name__ == '__main__':
    onboard_main()
