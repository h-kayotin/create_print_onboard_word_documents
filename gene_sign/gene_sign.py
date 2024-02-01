"""
gene_sign - 从excel读取数据，生产签名

logo图片：155px

Author: hanayo
Date： 2024/1/29
"""
import openpyxl
from openpyxl.worksheet.worksheet import Worksheet
from docx import Document
from docx.document import Document as Doc
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


class SignFTS:
    def __init__(self):
        workbook = openpyxl.load_workbook("签名生成.xlsx")
        worksheet = workbook['Sheet1']  # type: Worksheet
        # 获取excel中填写的内容
        self.sign_info_dict = dict()
        self.brand_info_list = list()
        for i in range(3, 13):
            self.sign_info_dict[worksheet[f"C{i}"].value] = str(worksheet[f"D{i}"].value).strip()
            if worksheet[f"H{i}"].value:
                self.brand_info_list.append(worksheet[f"G{i}"].value)
        # 根据图片数量，决定用哪个模板
        if self.brand_info_list:
            self.model = f"sign_modle_{len(self.brand_info_list)}.docx"
            self.file_name = f"生成签名_{self.sign_info_dict['name_cn']}_{len(self.brand_info_list)}.docx"
        else:
            self.model = "sign_modle_0.docx"
            self.file_name = f"生成签名_{self.sign_info_dict['name_cn']}_0.docx"

        print(self.sign_info_dict)
        print(self.brand_info_list)
        print(self.model)
        self.img_txt = ""
        for img in self.brand_info_list:
            self.img_txt += f"{img} | "
        self.img_txt = self.img_txt[:-2]

        self.doc = Document(f"models/{self.model}")  # type: Doc

        # 替换文字
        self.make_sign_word()
        # # 插入图片
        self.add_pic()

        self.doc.save(self.file_name)

    def make_sign_word(self):
        """替换模版文字"""
        for p in self.doc.paragraphs:
            print(p.text)
            for key, val in self.sign_info_dict.items():
                if key in p.text:
                    p.text = p.text.replace(key, val)
                    for run in p.runs:
                        run.font.size = Pt(12)
                        run.font.name = "微软雅黑"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 如果没有组，对这一行进行处理
        if self.sign_info_dict['group_cn'] == 'None':
            res_txt = self.doc.paragraphs[2].runs[0].text[12:]
            self.doc.paragraphs[2].runs[0].text = res_txt
            for run in self.doc.paragraphs[2].runs:
                run.font.size = Pt(12)
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_pic(self):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "img" in cell.text:
                        cell.text = cell.text.replace("img", "")
                        run = cell.paragraphs[0].runs[0]
                        run.add_picture(f"models/bran_images/{self.brand_info_list.pop(0)}.png", width=Cm(4))
                        if not self.brand_info_list:
                            return

        # for brand in self.brand_info_list:
        #     for p in self.doc.paragraphs:
        #         if brand in p.text:
        #             p.text = p.text.replace(brand, "")
        #             run = p.add_run()
        #             run.add_picture(f"models/bran_images/{brand}.png", width=Cm(3))


if __name__ == '__main__':
    my_s = SignFTS()
